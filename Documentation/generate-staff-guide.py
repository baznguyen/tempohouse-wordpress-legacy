"""
TEMPO House — Staff Content Guide Generator
Generates Documentation/staff-content-guide-v{VERSION}.docx

Run: python3 Documentation/generate-staff-guide.py
Requires: python-docx  (pip install python-docx)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Version ──────────────────────────────────────────────────────────────────
VERSION   = "1.0"
DOC_DATE  = "June 2026"
OUT_DIR   = os.path.dirname(os.path.abspath(__file__))
OUT_FILE  = os.path.join(OUT_DIR, f"staff-content-guide-v{VERSION}.docx")

# ── Brand palette ─────────────────────────────────────────────────────────────
INK        = RGBColor(0x1A, 0x18, 0x16)
TERRACOTTA = RGBColor(0xC7, 0x6E, 0x4B)
CREAM      = RGBColor(0xF7, 0xF3, 0xEE)
SAND       = RGBColor(0xE7, 0xD8, 0xC9)
SAGE       = RGBColor(0x8A, 0x92, 0x77)
AMBER      = RGBColor(0xDD, 0xAA, 0x62)
MUTED      = RGBColor(0x6B, 0x65, 0x60)

# System fonts — closest available matches to brand typefaces
FONT_DISPLAY = "Garamond"       # → Cormorant Garamond (editorial serif)
FONT_BODY    = "Calibri"        # → Space Grotesk (clean sans)
FONT_LABEL   = "Calibri"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell background with an RGBColor."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                    color="C76E4B", size=6):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_border_bottom(para, color="E7D8C9", size=6):
    """Add a bottom border rule to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shade(para, fill="F7F3EE"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)

def add_run(para, text, bold=False, italic=False, color=None,
            size=None, font=None, caps=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
    if caps:
        run.font.all_caps = True
    return run

def para_space(doc, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

# ── Document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins: generous (2.5cm sides, 3cm top/bottom)
for section in doc.sections:
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = FONT_BODY
style.font.size = Pt(10)
style.font.color.rgb = INK
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
style.paragraph_format.line_spacing = 1.3

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Terracotta rule at top
rule_top = doc.add_paragraph()
rule_top.paragraph_format.space_before = Pt(0)
rule_top.paragraph_format.space_after  = Pt(0)
set_para_shade(rule_top, fill="C76E4B")
rule_run = rule_top.add_run("  ")
rule_run.font.size = Pt(4)

# Spacer
for _ in range(8):
    doc.add_paragraph()

# TEMPO HOUSE label
eyebrow = doc.add_paragraph()
eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
eyebrow.paragraph_format.space_after = Pt(4)
add_run(eyebrow, "TEMPO HOUSE", bold=True, color=TERRACOTTA,
        size=8, font=FONT_BODY, caps=True)

# Document title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.paragraph_format.space_after = Pt(6)
add_run(title_p, "Website Management\nStaff Guide", color=INK,
        size=34, font=FONT_DISPLAY)

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
sub_p.paragraph_format.space_after = Pt(40)
add_run(sub_p, "How to create, manage, and publish\ncontent across the TEMPO House website.",
        color=MUTED, size=12, font=FONT_DISPLAY, italic=True)

# Version/date block — small info table
vtable = doc.add_table(rows=1, cols=3)
vtable.style = "Table Grid"
vtable.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["Version", "Date", "Prepared by"]
values = [f"v{VERSION}", DOC_DATE, "Raging Monk × TEMPO House"]
for i, cell in enumerate(vtable.rows[0].cells):
    set_cell_bg(cell, SAND)
    set_cell_border(cell, color="C76E4B", size=4)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = cell.add_paragraph()
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(0)
    add_run(lp, labels[i].upper() + "\n", bold=True, color=MUTED, size=7, font=FONT_BODY, caps=True)
    add_run(lp, values[i], color=INK, size=9, font=FONT_BODY)

for _ in range(6):
    doc.add_paragraph()

# Bottom terracotta rule
rule_bot = doc.add_paragraph()
rule_bot.paragraph_format.space_before = Pt(0)
rule_bot.paragraph_format.space_after  = Pt(2)
set_para_shade(rule_bot, fill="C76E4B")
rb_run = rule_bot.add_run("  ")
rb_run.font.size = Pt(4)

# Bottom note
note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
note_p.paragraph_format.space_after = Pt(0)
add_run(note_p, "Internal use only — tempohouse.com.vn",
        color=MUTED, size=8, font=FONT_BODY)

# Page break
doc.add_page_break()

# ── Section heading helper ────────────────────────────────────────────────────
def h1(doc, text):
    """Major section heading — terracotta, Garamond, ruled."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    set_para_border_bottom(p, color="C76E4B", size=8)
    add_run(p, text, color=TERRACOTTA, size=20, font=FONT_DISPLAY)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, color=INK, size=14, font=FONT_DISPLAY)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text.upper(), bold=True, color=MUTED, size=8,
            font=FONT_BODY, caps=True)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.5) if indent else Cm(0)
    add_run(p, text, color=INK, size=10, font=FONT_BODY)
    return p

def bullet(doc, text, sub=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(1.2 if sub else 0.5)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text, color=INK, size=10, font=FONT_BODY)
    return p

def callout(doc, label, text, bg="F7F3EE", border="C76E4B"):
    """Highlighted callout box via a 1-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    set_cell_bg(c, RGBColor(
        int(bg[0:2], 16), int(bg[2:4], 16), int(bg[4:6], 16)))
    set_cell_border(c, sides=("left",), color=border, size=18)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    if label:
        add_run(p, label.upper() + "  ", bold=True, color=TERRACOTTA,
                size=8, font=FONT_BODY)
    add_run(p, text, color=INK, size=10, font=FONT_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def field_table(doc, rows_data):
    """
    Two-column reference table: Field name | Description.
    rows_data: list of (field_label, description) tuples.
    """
    t = doc.add_table(rows=len(rows_data) + 1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, label in enumerate(["Field", "What to put here"]):
        c = hdr.cells[i]
        set_cell_bg(c, INK)
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        add_run(p, label.upper(), bold=True, color=CREAM, size=8, font=FONT_BODY, caps=True)
    # Data rows
    for ri, (fname, fdesc) in enumerate(rows_data):
        row = t.rows[ri + 1]
        bg  = SAND if ri % 2 == 0 else CREAM
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        p0 = c0.add_paragraph()
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after  = Pt(3)
        add_run(p0, fname, bold=True, color=INK, size=9, font=FONT_BODY)
        p1 = c1.add_paragraph()
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after  = Pt(3)
        add_run(p1, fdesc, color=INK, size=9, font=FONT_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════

toc_title = doc.add_paragraph()
toc_title.paragraph_format.space_before = Pt(0)
toc_title.paragraph_format.space_after  = Pt(8)
add_run(toc_title, "Contents", color=INK, size=18, font=FONT_DISPLAY)

toc_items = [
    ("1", "Overview"),
    ("2", "Managing Events — The Core Workflow"),
    ("3", "Creating an Event Post — Step by Step"),
    ("4", "Event Fields Reference"),
    ("5", "Managing Event Status"),
    ("6", "The What's On Page"),
    ("7", "Brand Voice & Writing Guidelines"),
    ("8", "Image & Artwork Guidelines"),
    ("9", "SEO Best Practices"),
    ("10", "Frequently Asked Questions"),
    ("11", "Version History"),
]
for num, item in toc_items:
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(3)
    add_run(tp, f"{num}.  ", bold=True, color=TERRACOTTA, size=10, font=FONT_BODY)
    add_run(tp, item, color=INK, size=10, font=FONT_BODY)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "1  Overview")

body(doc, "This guide covers everything your team needs to manage content on the TEMPO House website. You don't need technical knowledge — if you can write a blog post, you can manage events, update copy, and keep the programme current.")

body(doc, "The website runs on WordPress, accessed at your WordPress admin dashboard. Events, copy, and media are all managed from there. The site updates in real time — changes are live the moment you publish.")

callout(doc, "Key principle",
        "The 'event' and 'active' tags are what drive the What's On page and the homepage carousel. "
        "No event appears anywhere on the site unless it has the 'event' tag. "
        "Add 'active' when the event is live. Remove it when it's over.")

h2(doc, "Who this guide is for")
bullet(doc, "Marketing coordinators managing the events programme")
bullet(doc, "Anyone creating or updating event posts")
bullet(doc, "Team leads reviewing the site before it goes live")

h2(doc, "What you can manage")
bullet(doc, "Events — create, publish, activate, archive")
bullet(doc, "Event copy — titles, descriptions, dates, pricing")
bullet(doc, "Event imagery — poster photos and card artwork")
bullet(doc, "What's On page — which events appear in which section")
bullet(doc, "Homepage carousel — automatically reflects active events")

# ═══════════════════════════════════════════════════════════════════════════════
# 2. THE CORE WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "2  Managing Events — The Core Workflow")

body(doc, "All events on the TEMPO House website are standard WordPress Posts. They are distinguished from regular posts by having the tag 'event'. The tag 'active' controls whether an event appears under 'Happening Now' on the What's On page and in the homepage carousel.")

h2(doc, "The event tag system")

# Status table
st = doc.add_table(rows=4, cols=3)
st.style = "Table Grid"
st.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ["Tags on the post", "Where it appears", "When to use"]
for i, h in enumerate(headers):
    c = st.rows[0].cells[i]
    set_cell_bg(c, INK)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, h.upper(), bold=True, color=CREAM, size=8, font=FONT_BODY, caps=True)

status_rows = [
    ("event  +  active",   "Homepage carousel  ·  What's On → Happening Now",
     "Event is currently running or ongoing"),
    ("event  (no active)", "What's On → Coming Up only",
     "Event is announced but not yet started"),
    ("Neither tag",        "Not visible on site",
     "Draft / archived / non-event post"),
]
bgs = [SAND, CREAM, SAND]
for ri, (tags, where, when) in enumerate(status_rows):
    row = st.rows[ri + 1]
    for ci, text in enumerate([tags, where, when]):
        c = row.cells[ci]
        set_cell_bg(c, bgs[ri])
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        bold = ci == 0
        add_run(p, text, bold=bold, color=INK, size=9, font=FONT_BODY)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

h2(doc, "Lifecycle of a typical event")
steps = [
    ("1  Create", "Write the post in draft — fill in all Event Details fields, add copy, upload poster."),
    ("2  Publish", "Set status to Published. Add the tag 'event'. It appears on What's On under 'Coming Up'."),
    ("3  Activate", "When the event starts, add the tag 'active'. It moves to 'Happening Now' and the homepage."),
    ("4  Archive", "When the event ends, remove the 'active' tag (it leaves 'Happening Now'). "
                   "Then either remove 'event' tag or unpublish to hide it entirely."),
]
for num, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, num + "  ", bold=True, color=TERRACOTTA, size=10, font=FONT_BODY)
    add_run(p, desc, color=INK, size=10, font=FONT_BODY)

callout(doc, "Tip",
        "You don't need to unpublish old events to remove them from the site. "
        "Simply removing the 'active' tag drops it out of 'Happening Now' and the homepage. "
        "Removing the 'event' tag hides it from the What's On page entirely — "
        "but the post URL still works and is indexed by Google.")

# ═══════════════════════════════════════════════════════════════════════════════
# 3. STEP-BY-STEP: CREATING AN EVENT POST
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "3  Creating an Event Post — Step by Step")

steps_detail = [
    ("Step 1", "Log in to WordPress Admin",
     "Go to your WordPress admin URL. Log in with your credentials."),
    ("Step 2", "Go to Posts → Add New",
     "In the left sidebar, click Posts, then Add New. This opens the Gutenberg editor."),
    ("Step 3", "Write the title",
     "The title is the event name — exactly as you want it to appear on the card and detail page. "
     "Keep it short and direct. Examples: 'TEMPO Sessions', 'Cocktail Masterclass', 'Opening Night — Mai Nguyen'."),
    ("Step 4", "Write the body content",
     "The body is the full event description — shown on the event detail page. Use Gutenberg blocks: "
     "paragraphs for main copy, headings for sections (H3 only, not H2), bullet lists for practical details "
     "(what to expect, what's included, booking info). "
     "Write 3–5 paragraphs. See Section 7 for brand voice guidelines."),
    ("Step 5", "Fill in Event Details (right sidebar)",
     "Scroll down in the right sidebar to find the Event Details panel. "
     "Fill in every field that applies. See Section 4 for a full field reference."),
    ("Step 6", "Add the excerpt",
     "In the right sidebar under 'Excerpt', write a 1–2 sentence summary of the event. "
     "This is used as the meta description for Google and social sharing. "
     "Keep it under 160 characters."),
    ("Step 7", "Set the Featured Image",
     "The Featured Image is used as a fallback if no Poster Image is set in Event Details. "
     "Upload a portrait or square image. Minimum 800px wide."),
    ("Step 8", "Add tags",
     "In the Tags field (right sidebar, under the Category panel), type 'event' and press Enter. "
     "If the event is live right now, also add 'active'. "
     "Tags are comma-separated — you can add both at once: event, active"),
    ("Step 9", "Publish",
     "Click Publish (top right). The event is now live. "
     "If you tagged it 'active', it appears on the homepage immediately."),
]

for step, title, detail in steps_detail:
    h3(doc, step + "  —  " + title)
    body(doc, detail)

callout(doc, "Important",
        "Do not change the post Category — leave it as 'Uncategorized'. "
        "Event categorisation is handled by the Event Details fields, not WordPress categories. "
        "The tag system ('event', 'active') is what controls visibility on the site.")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. EVENT FIELDS REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "4  Event Fields Reference")

body(doc, "These fields appear in the Event Details panel in the right sidebar of the post editor. "
     "They control how the event appears on the What's On page and the detail page.")

field_table(doc, [
    ("Category",
     "The event type label. Shown on the card and detail page. "
     "Examples: Live Music, Exhibition, Workshop, Private Dining, Gallery Launch, Brand Event."),
    ("Event Date",
     "The date the event runs. Leave blank for recurring events (monthly sessions, ongoing exhibitions). "
     "Format: pick from the date picker — displays as 'Saturday 12 July 2026'."),
    ("End Date",
     "For multi-day events only — the last day the event runs. Leave blank for single-day events."),
    ("Time",
     "Display time string shown on the card and meta bar. "
     "Format: '20:00 – 23:00'. Use an en dash (–), not a hyphen (-)."),
    ("Recurrence",
     "How often the event repeats. Options: One Night Only, Weekly, Monthly, Ongoing. "
     "Leave blank for unique dated events."),
    ("Card Colour",
     "The background colour of the event card. Choose: Dark Ink (night/music), "
     "Warm Sand (gallery/art), Soft Cream (daytime/workshop)."),
    ("Poster Image",
     "The large hero image shown at the top of the event detail page. "
     "Landscape ratio, minimum 1920×1080px. Upload event artwork, photography, or a designed poster."),
    ("Admission",
     "Shown in the meta bar on the detail page. Examples: Free, Free entry, 150,000 VND, "
     "Ticketed, From 200,000 VND."),
    ("Ticket / RSVP Link",
     "Optional. If provided, a 'Get Tickets' button appears on the detail page. "
     "Paste the full URL (https://...)."),
    ("Card Artwork Type",
     "Controls what appears inside the event card on the What's On page. "
     "None = shows Category text (recommended for most events). "
     "Image = shows a portrait photo. Video = plays an MP4 clip."),
    ("Card Artwork",
     "Only shown when Card Artwork Type is set to Image or Video. "
     "Portrait image: 600×800px recommended. This is what's seen on the card itself."),
])

callout(doc, "Card Colour guide",
        "Dark Ink → music nights, late-night events, cocktail events.\n"
        "Warm Sand → art exhibitions, gallery launches, creative workshops.\n"
        "Soft Cream → daytime events, café sessions, morning masterclasses.")

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MANAGING EVENT STATUS
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "5  Managing Event Status")

h2(doc, "Activating an event (moving to 'Happening Now')")
bullet(doc, "Open the event post in WordPress Admin → Posts")
bullet(doc, "In the Tags field, add the tag 'active'")
bullet(doc, "Click Update (top right)")
bullet(doc, "The event immediately moves to the 'Happening Now' section on the What's On page and appears in the homepage carousel")

h2(doc, "Deactivating an event (removing from 'Happening Now')")
bullet(doc, "Open the event post")
bullet(doc, "Click the ✕ next to the 'active' tag to remove it")
bullet(doc, "Click Update")
bullet(doc, "The event moves back to 'Coming Up' (if it has a future date) or disappears from What's On")

h2(doc, "Hiding an event entirely")
bullet(doc, "Remove both the 'event' and 'active' tags — the post is no longer visible on any listing page")
bullet(doc, "OR set the post status to 'Draft' — this also removes it from Google")

h2(doc, "Changing event details")
bullet(doc, "Open the post, make edits to any field or body content, click Update")
bullet(doc, "Changes are live immediately — no cache to clear")

callout(doc, "Note",
        "Editing a published post does not affect its URL or SEO ranking. "
        "Changing the post title after publishing will NOT change the URL (WordPress locks the slug on first publish). "
        "If you need to change the URL, contact your web manager.")

# ═══════════════════════════════════════════════════════════════════════════════
# 6. THE WHAT'S ON PAGE
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "6  The What's On Page")

body(doc, "The What's On page at tempohouse.com.vn/whats-on/ is automatically generated from your event posts. "
     "You do not edit it directly — it updates in real time as you publish and tag events.")

h2(doc, "Page structure")
bullet(doc, "Happening Now — all posts tagged 'event' AND 'active'. Sorted newest published first.")
bullet(doc, "Coming Up — all posts tagged 'event' but NOT 'active', with a future date or no date set.")
bullet(doc, "If neither section has content, the page shows placeholder cards until real events are published.")

h2(doc, "Homepage carousel")
bullet(doc, "Shows the 3 most recently published posts tagged 'event' AND 'active'.")
bullet(doc, "To change which events appear: publish new active events, or adjust publish dates.")
bullet(doc, "To show a specific event first: set its publish date to the most recent.")

callout(doc, "Maximum on homepage",
        "The homepage carousel always shows exactly 3 events. "
        "If more than 3 are tagged 'active', only the 3 most recently published appear on the homepage "
        "(all of them appear on the What's On page).")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. BRAND VOICE & WRITING GUIDELINES
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "7  Brand Voice & Writing Guidelines")

body(doc, "TEMPO House has a specific voice: unhurried, specific, quietly confident. "
     "The writing avoids hype. It earns attention through precision, not volume.")

h2(doc, "The three modes")
modes = [
    ("Day (café)",
     "Direct and grounded. The morning crowd knows what they're here for. "
     "Pull up a stool or claim the corner."),
    ("Night (bar)",
     "Relaxed, a little cinematic. The lights don't change dramatically. "
     "The room just shifts. You'll notice it around the second drink."),
    ("Gallery",
     "Understated, respectful of the work. Level 2 is open for the current exhibition. "
     "There's no guided tour. Stay as long as the work asks you to."),
]
for mode, example in modes:
    h3(doc, mode)
    body(doc, "Example: " + example, indent=True)

h2(doc, "Do")
bullet(doc, "Write in short, declarative sentences")
bullet(doc, "Name the specific: the act, the cocktail, the artist")
bullet(doc, "Use precise times and dates — '20:00 – 23:00', not 'evening'")
bullet(doc, "Let the details carry the energy — no need to tell people it will be 'amazing'")
bullet(doc, "Keep the excerpt under 160 characters — one or two sentences maximum")

h2(doc, "Don't")
bullet(doc, "Use: curated, artisanal, vibe, amazing, awesome, immersive, unforgettable")
bullet(doc, "Use 'experience' as a noun (as in 'a unique experience')")
bullet(doc, "Write in the third person ('TEMPO House is proud to present...')")
bullet(doc, "Use exclamation marks")
bullet(doc, "Pad with filler — if a sentence doesn't add information, cut it")

callout(doc, "Test",
        "Read the copy out loud. If it sounds like marketing, rewrite it to sound like a person.")

# ═══════════════════════════════════════════════════════════════════════════════
# 8. IMAGE & ARTWORK GUIDELINES
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "8  Image & Artwork Guidelines")

h2(doc, "Poster Image (event detail page hero)")
bullet(doc, "Used as: full-bleed hero at the top of the event detail page")
bullet(doc, "Dimensions: 1920 × 1080px minimum (landscape/16:9)")
bullet(doc, "Format: JPG or PNG, under 2MB")
bullet(doc, "Content: event artwork, photography, designed poster, or artist work")
bullet(doc, "If not set, the hero shows a coloured background based on the Card Colour setting")

h2(doc, "Card Artwork (What's On card)")
bullet(doc, "Used as: background image inside the card on the What's On page and homepage")
bullet(doc, "Dimensions: 600 × 800px (portrait/3:4)")
bullet(doc, "Format: JPG, under 1MB")
bullet(doc, "If not set, the card shows large category text (e.g. 'Live Music') as a graphic element — this is the intended design and works well")

h2(doc, "Featured Image (fallback)")
bullet(doc, "Used as: fallback if no Poster Image is set")
bullet(doc, "Set via 'Featured Image' in the right sidebar")
bullet(doc, "Minimum 800px wide")

h2(doc, "General image rules")
bullet(doc, "Upload the highest quality image you have — WordPress generates all crop sizes automatically")
bullet(doc, "Avoid images with logos, watermarks, or text overlaid — they clash with the site's typography")
bullet(doc, "Faces, hands, and close-up texture shots perform better than wide empty-room shots")
bullet(doc, "Black and white images work particularly well with the Dark Ink card colour")

callout(doc, "No image yet?",
        "Don't let missing artwork delay publishing. Set the Card Colour and publish — "
        "the category ghost text looks intentional and designed. "
        "Upload the poster when the artwork is ready.")

# ═══════════════════════════════════════════════════════════════════════════════
# 9. SEO BEST PRACTICES
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "9  SEO Best Practices")

body(doc, "The website handles most SEO automatically. Event posts get structured data (Schema.org Event markup) "
     "generated from the Event Details fields — this helps Google show your events in search results with dates and times.")

h2(doc, "What you control")
field_table(doc, [
    ("Post title",
     "The H1 of the page and the main SEO signal. Be specific: "
     "'TEMPO Sessions — June' outperforms 'Live Music Event'."),
    ("Excerpt",
     "This becomes the meta description shown in Google search results. "
     "Write it as a search result snippet: include what, when, where. Under 160 characters."),
    ("Event Date + Time",
     "Used in the auto-generated Schema markup. Always fill these in for one-off events — "
     "Google uses them to show event details in search results."),
    ("Poster Image",
     "Used as the Open Graph image for social sharing (Facebook, Zalo, iMessage previews). "
     "A strong image significantly improves click-through from social links."),
    ("Admission",
     "Included in Schema.org offer markup — can appear in Google's event rich results."),
])

h2(doc, "URL structure")
body(doc, "Event URLs are automatically generated from the post title: "
     "tempohouse.com.vn/tempo-sessions/. "
     "The URL is set when you first publish — changing the title after publishing does not change the URL. "
     "Choose a clear, concise title before publishing.")

callout(doc, "Don't overthink it",
        "Write clearly for the reader first. The SEO fields (excerpt, date, image) "
        "take 2 minutes to fill in and do most of the work automatically.")

# ═══════════════════════════════════════════════════════════════════════════════
# 10. FAQ
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "10  Frequently Asked Questions")

faqs = [
    ("I added the 'active' tag but the event isn't on the homepage.",
     "Check that you also have the 'event' tag on the post — both are required for the homepage. "
     "Click Update after adding the tag."),
    ("The What's On page shows placeholder cards instead of my events.",
     "Make sure your events are Published (not Draft) and have the 'event' tag. "
     "The placeholder cards only show when no published event-tagged posts exist."),
    ("I want to change the order of events on the What's On page.",
     "The 'Happening Now' section sorts by publish date, newest first. "
     "To move an event to the top, edit it and change its publish date to today — "
     "click the date shown under 'Publish' in the sidebar and update it."),
    ("I accidentally deleted an event post.",
     "Go to Posts → Trash in WordPress admin. Find the post and click Restore. "
     "If it's not there, contact your web manager — posts in Trash are kept for 30 days."),
    ("Can I schedule an event post to publish in the future?",
     "Yes. In the Publish panel (right sidebar), click the date shown next to 'Publish' "
     "and set a future date and time. The post will publish automatically. "
     "Add the 'event' tag before scheduling so it appears on the What's On page when it goes live."),
    ("Who manages SEO optimisations and ad campaign assets?",
     "SEO optimisations (meta descriptions, Schema markup, keyword analysis) "
     "are reviewed by the Raging Monk team. For ad campaign assets and social media copy, "
     "raise a request via your usual channel — the Muse AI system can generate drafts "
     "from your event details."),
    ("How do I add a Vietnamese version of an event?",
     "For now, write bilingual copy within the same post body — Vietnamese first, then English "
     "(or vice versa, separated by a horizontal rule block). "
     "A full bilingual content system is planned for a future version of the site."),
]

for q, a in faqs:
    h3(doc, q)
    body(doc, a)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════════════════════════════════════════
# 11. VERSION HISTORY
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "11  Version History")

vt = doc.add_table(rows=2, cols=4)
vt.style = "Table Grid"
vt.alignment = WD_TABLE_ALIGNMENT.LEFT
vh = ["Version", "Date", "Author", "Changes"]
for i, label in enumerate(vh):
    c = vt.rows[0].cells[i]
    set_cell_bg(c, INK)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, label.upper(), bold=True, color=CREAM, size=8, font=FONT_BODY, caps=True)

v1_data = [f"v{VERSION}", DOC_DATE, "Raging Monk",
           "Initial release — events workflow, field reference, brand voice, SEO, FAQ."]
for ci, val in enumerate(v1_data):
    c = vt.rows[1].cells[ci]
    set_cell_bg(c, SAND)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, val, color=INK, size=9, font=FONT_BODY)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

callout(doc, "Updates",
        f"This document is versioned. When processes change, a new version (v1.1, v2.0, etc.) "
        f"will be issued with a changelog entry above. Always refer to the latest version — "
        f"check the filename for the version number.",
        bg="F0EBE3", border="8A9277")

# ── Footer line ───────────────────────────────────────────────────────────────
fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(20)
fp.paragraph_format.space_after  = Pt(0)
set_para_border_bottom(fp, color="E7D8C9", size=6)
add_run(fp, f"TEMPO House — Staff Content Guide  v{VERSION}  ·  {DOC_DATE}  ·  Internal use only",
        color=MUTED, size=8, font=FONT_BODY)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"✓  Saved: {OUT_FILE}")
